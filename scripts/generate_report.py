import json
from docx import Document

def generate_report(data):
    doc = Document()
    #lang = data["lang"]
    lang = "FR"
    
    def add_heading(title):
        doc.add_heading(title, level=1)

    def add_paragraph(text):
        doc.add_paragraph(text)

    # Section: Extraction
    add_heading("Extraction")
    if lang == "FR":
        text = "L’ADN de [...] a été extrait avec le kit [...] (réf. [...] )."
    else:
        text = "DNA from [...]  was extracted using the [...] kit (Ref. [...] )."
    add_paragraph(text)

    # Section: PCR
    add_heading("PCR")
    if lang == "FR":
        text = " PCR ont été réalisées pour amplifier un fragment de [...]."
        text += " Les PCR des échantillons  ont échoué, mais ont tout de même été séquencées."
    else:
        text = "PCRs were performed to amplify a {data['amplicon']} fragment."
    add_paragraph(text)

    # Section: Séquençage
    add_heading("Séquençage" if lang == "FR" else "Sequencing")
    text = f"Séquençage effectué sur MinION MK1C avec des Flongles (voir [...] )." if lang == "FR" else f"Sequencing performed on )."
    add_paragraph(text)

    # Section: Bioinformatique
    add_heading("Analyses bioinformatiques" if lang == "FR" else "Bioinformatics")
    if lang == "FR":
        add_paragraph("Les analyses bio-informatiques ont été réalisées:")
        add_paragraph(
            "- par BIG (plateau de bio-informatique de PlantBios https://institut-sophia-agrobiotech.paca.hub.inrae.fr/infrastructure-plantbios/plateau-de-bioinformatique-et-genomique) pour la concaténation des fichiers et le basecalling « Super High Accuracy »."
        )
        add_paragraph(
            "- par le pipeline bioinformatique OpenMetaBar avec :"
            "- minibar pour :\n"
            "  • le démultiplexage : séparation des lectures et regroupement par individus extraits ;\n"
        )
        add_paragraph(
            "- seqkit pour :\n"
            "  • le filtrage de qualité : prise en compte de la qualité, de la longueur et du contenu attendu (séquences codantes sans codon stop) ;\n"
        )
        add_paragraph(
            "- Lotus3 pour la création de cluster, génération de Zotu et blast des séquences."
        )


    # Section: BLAST
    add_heading("BLAST")
    add_paragraph("blast_result")

    # Section: Clustering
    add_heading("Clustering")
    add_paragraph("clustering_result")

    # Section: Conclusion
    add_heading("Conclusion")
    add_paragraph("conclusion")

    # Citation IDMABIO
    add_heading("Remerciements" if lang == "FR" else "Acknowledgements")
    if lang == "FR":
        citation = (
            "L'identification moléculaire a été réalisée sur la plateforme IDMABIO (https://idmabio.com).\n"
            "Nous remercions la plateforme IDMABIO (ISC PlantBIOs, https://doi.org/10.15454/qyey-ar89) pour les analyses moléculaires."
        )
    else:
        citation = (
            "Molecular identification was performed using the IDMABIO platform (https://idmabio.com).\n"
            "We are grateful to IDMABIO (ISC PlantBIOs, https://doi.org/10.15454/qyey-ar89) for molecular identification and analysis."
        )
    add_paragraph(citation)

    # Traçabilité
    add_heading("Traçabilité")
    add_paragraph("Document généré automatiquement. Localisation: ")

    return doc

# === MAIN ===
if __name__ == "__main__":
    #with open("generate_reports/info.json", "r", encoding="utf-8") as f:
        #data = json.load(f)

    doc = generate_report("test")
    doc.save("rapport_final.docx")